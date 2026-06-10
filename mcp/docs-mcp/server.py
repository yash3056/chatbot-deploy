#!/usr/bin/env python3
"""
Docs MCP Server — PDF reading and DOCX read/write over SSE/HTTP.

Tools
─────
  pdf_read   Extract text from a PDF (full doc or page range, with table support)
  pdf_info   Metadata + structure info for a PDF
  docx_read  Read a .docx file and return clean Markdown
  docx_write Create or update a .docx file from Markdown content

Transport: SSE (Server-Sent Events) on port 8099
  /sse   — MCP SSE endpoint (LibreChat / MCP clients)
  /help  — Human-readable tool listing
"""

import json
import os
import re
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP
from starlette.requests import Request
from starlette.responses import HTMLResponse
from starlette.routing import Route

# ──────────────────────────────────────────────────────────────────────────────
# Config
# ──────────────────────────────────────────────────────────────────────────────

PORT = int(os.getenv("DOCS_MCP_PORT", "8099"))

mcp = FastMCP("docs-mcp", host="0.0.0.0", port=PORT)


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def _resolve(path: str) -> Path:
    """Resolve path; raise FileNotFoundError if file does not exist."""
    p = Path(path).expanduser().resolve()
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    return p


def _parse_page_range(spec: str, total: int) -> list[int]:
    """
    Convert a page-range spec string to a sorted list of 0-based page indices.
    Spec examples: '1', '1-5', '2,4,6', '3-', '-5', '1-3,7,9-11'
    """
    indices: set[int] = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            lo_s, hi_s = part.split("-", 1)
            lo = int(lo_s) - 1 if lo_s.strip() else 0
            hi = int(hi_s) - 1 if hi_s.strip() else total - 1
            indices.update(range(max(0, lo), min(total, hi + 1)))
        else:
            idx = int(part) - 1
            if 0 <= idx < total:
                indices.add(idx)
    return sorted(indices)


# ──────────────────────────────────────────────────────────────────────────────
# Markdown → DOCX converter
# ──────────────────────────────────────────────────────────────────────────────

def _md_to_docx(content: str, path: str, title: str | None, author: str | None, page_size: str) -> str:
    """Parse Markdown and write a .docx via python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # page size
    section = doc.sections[0]
    if page_size == "Letter":
        section.page_width  = int(8.5 * 914400)
        section.page_height = int(11  * 914400)
    else:  # A4
        section.page_width  = int(21  * 914400 / 2.54)
        section.page_height = int(29.7 * 914400 / 2.54)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # core properties
    props = doc.core_properties
    if title:  props.title  = title
    if author: props.author = author

    # inline formatting
    _INLINE_RE = re.compile(
        r"\*\*\*(.+?)\*\*\*"    # bold+italic
        r"|\*\*(.+?)\*\*"       # bold
        r"|\*(.+?)\*"           # italic
        r"|`(.+?)`"             # inline code
        r"|\[(.+?)\]\((.+?)\)"  # hyperlink
    )

    def _add_inline(para, text: str):
        pos = 0
        for m in _INLINE_RE.finditer(text):
            if m.start() > pos:
                para.add_run(text[pos:m.start()])
            if m.group(1):
                r = para.add_run(m.group(1)); r.bold = True; r.italic = True
            elif m.group(2):
                r = para.add_run(m.group(2)); r.bold = True
            elif m.group(3):
                r = para.add_run(m.group(3)); r.italic = True
            elif m.group(4):
                r = para.add_run(m.group(4))
                r.font.name = "Courier New"; r.font.size = Pt(10)
            elif m.group(5):
                _add_hyperlink(para, m.group(6), m.group(5))
            pos = m.end()
        if pos < len(text):
            para.add_run(text[pos:])

    def _add_hyperlink(paragraph, url: str, text: str):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        new_run_text = OxmlElement("w:t")
        new_run_text.text = text
        new_run.append(new_run_text)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _add_table(rows: list[list[str]], has_header: bool):
        if not rows:
            return
        cols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Table Grid"
        for r_idx, row_cells in enumerate(rows):
            for c_idx in range(cols):
                cell = tbl.cell(r_idx, c_idx)
                cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                para = cell.paragraphs[0]
                if r_idx == 0 and has_header:
                    run = para.add_run(cell_text.strip())
                    run.bold = True
                else:
                    _add_inline(para, cell_text.strip())
        doc.add_paragraph()

    def _is_table_row(line: str) -> bool:
        return line.strip().startswith("|")

    def _parse_table_row(line: str) -> list[str]:
        stripped = line.strip().strip("|")
        return [c for c in stripped.split("|")]

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # table block
        if _is_table_row(stripped):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            rows = []
            has_header = False
            for idx, tl in enumerate(table_lines):
                if re.match(r"^\|?[\s:\-|]+\|?$", tl.strip()):
                    if idx == 1:
                        has_header = True
                    continue
                rows.append(_parse_table_row(tl))
            _add_table(rows, has_header)
            continue

        # unordered list
        if re.match(r"^[-*+]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, re.sub(r"^[-*+]\s+", "", stripped))
            i += 1
            continue

        # ordered list
        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            for cl in code_lines:
                para = doc.add_paragraph(style="No Spacing")
                run = para.add_run(cl)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            doc.add_paragraph()
            continue

        # blockquote
        if stripped.startswith("> "):
            style_names = [s.name for s in doc.styles]
            para = doc.add_paragraph(style="Quote") if "Quote" in style_names else doc.add_paragraph()
            _add_inline(para, stripped[2:])
            para.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue

        # plain paragraph (collect continuation lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nl = lines[i].strip()
            if (not nl or nl.startswith("#") or nl.startswith("```")
                    or re.match(r"^[-*+]\s+", nl) or re.match(r"^\d+\.\s+", nl)
                    or _is_table_row(nl) or re.match(r"^(-{3,}|\*{3,}|_{3,})$", nl)):
                break
            para_lines.append(nl)
            i += 1
        para = doc.add_paragraph()
        _add_inline(para, " ".join(para_lines))

    out = Path(path).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ──────────────────────────────────────────────────────────────────────────────
# MCP Tools
# ──────────────────────────────────────────────────────────────────────────────

@mcp.tool()
def pdf_read(
    path: str,
    pages: str = "",
    extract_tables: bool = False,
    combine: bool = True,
) -> str:
    """
    Extract text from a PDF file.

    Args:
        path: Absolute or relative path to the PDF file.
        pages: Page range to extract, e.g. '1', '1-5', '2,4,6', '3-'. 1-indexed. Omit to extract all.
        extract_tables: Also extract tables from each page (default: false).
        combine: Include a single combined text string in addition to per-page results (default: true).

    Returns:
        JSON with per-page text, char counts, optional tables, and combined text.
    """
    try:
        import pdfplumber

        pdf_path = _resolve(path)
        results = []
        combined_parts = []

        with pdfplumber.open(str(pdf_path)) as pdf:
            total = len(pdf.pages)
            indices = (
                _parse_page_range(pages, total)
                if pages else list(range(total))
            )

            for idx in indices:
                page = pdf.pages[idx]
                text = page.extract_text() or ""
                entry: dict[str, Any] = {
                    "page": idx + 1,
                    "text": text,
                    "char_count": len(text),
                }
                if extract_tables:
                    raw_tables = page.extract_tables()
                    entry["tables"] = raw_tables if raw_tables else []
                results.append(entry)
                combined_parts.append(f"--- Page {idx + 1} ---\n{text}")

        out: dict[str, Any] = {
            "file": str(pdf_path),
            "total_pages": total,
            "extracted_pages": len(results),
            "pages": results,
        }
        if combine:
            out["combined_text"] = "\n\n".join(combined_parts)

        return json.dumps(out, ensure_ascii=False, indent=2)

    except FileNotFoundError as e:
        return json.dumps({"error": str(e)})
    except Exception as e:
        import traceback
        return json.dumps({"error": f"{type(e).__name__}: {e}\n{traceback.format_exc()}"})


@mcp.tool()
def pdf_info(path: str) -> str:
    """
    Return metadata and structural information for a PDF.

    Args:
        path: Absolute or relative path to the PDF file.

    Returns:
        JSON with title, author, page count, page sizes, encryption status, and dates.
    """
    try:
        from pypdf import PdfReader

        pdf_path = _resolve(path)
        reader = PdfReader(str(pdf_path))
        meta = reader.metadata or {}

        def _clean(v):
            if v is None:
                return None
            s = str(v)
            m = re.match(r"D:(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})", s)
            if m:
                return f"{m[1]}-{m[2]}-{m[3]} {m[4]}:{m[5]}:{m[6]}"
            return s

        page_sizes = []
        for i, page in enumerate(reader.pages):
            w = float(page.mediabox.width)
            h = float(page.mediabox.height)
            page_sizes.append({
                "page": i + 1,
                "width_pt": round(w, 2),
                "height_pt": round(h, 2),
                "width_in": round(w / 72, 3),
                "height_in": round(h / 72, 3),
            })

        unique_sizes = {(s["width_pt"], s["height_pt"]) for s in page_sizes}

        return json.dumps({
            "file": str(pdf_path),
            "file_size_bytes": pdf_path.stat().st_size,
            "page_count": len(reader.pages),
            "encrypted": reader.is_encrypted,
            "metadata": {
                "title":    _clean(meta.get("/Title")),
                "author":   _clean(meta.get("/Author")),
                "subject":  _clean(meta.get("/Subject")),
                "keywords": _clean(meta.get("/Keywords")),
                "creator":  _clean(meta.get("/Creator")),
                "producer": _clean(meta.get("/Producer")),
                "created":  _clean(meta.get("/CreationDate")),
                "modified": _clean(meta.get("/ModDate")),
            },
            "page_sizes": page_sizes,
            "unique_page_sizes": [
                {"width_pt": w, "height_pt": h,
                 "width_in": round(w/72, 3), "height_in": round(h/72, 3)}
                for w, h in sorted(unique_sizes)
            ],
        }, ensure_ascii=False, indent=2)

    except FileNotFoundError as e:
        return json.dumps({"error": str(e)})
    except Exception as e:
        import traceback
        return json.dumps({"error": f"{type(e).__name__}: {e}\n{traceback.format_exc()}"})


@mcp.tool()
def docx_read(path: str, include_messages: bool = False) -> str:
    """
    Read a .docx Word document and return its content as clean Markdown.
    Preserves headings, bold, italic, lists, tables, and hyperlinks.

    Args:
        path: Absolute or relative path to the .docx file.
        include_messages: Include mammoth conversion warnings/messages (default: false).

    Returns:
        JSON with markdown content, file size, and optional conversion messages.
    """
    try:
        import mammoth

        docx_path = _resolve(path)

        with open(str(docx_path), "rb") as f:
            result = mammoth.convert_to_markdown(f)

        out: dict[str, Any] = {
            "file": str(docx_path),
            "file_size_bytes": docx_path.stat().st_size,
            "markdown": result.value,
            "char_count": len(result.value),
        }
        if include_messages and result.messages:
            out["messages"] = [str(m) for m in result.messages]

        return json.dumps(out, ensure_ascii=False, indent=2)

    except FileNotFoundError as e:
        return json.dumps({"error": str(e)})
    except Exception as e:
        import traceback
        return json.dumps({"error": f"{type(e).__name__}: {e}\n{traceback.format_exc()}"})


@mcp.tool()
def docx_write(
    path: str,
    content: str,
    title: str = "",
    author: str = "",
    page_size: str = "A4",
) -> str:
    """
    Create a new .docx file (or overwrite an existing one) from Markdown content.
    Supports headings, bold, italic, lists, tables, code blocks, blockquotes, and hyperlinks.

    Args:
        path: Destination .docx file path (created or overwritten).
        content: Markdown text to write into the document.
        title: Document title stored in core properties (optional).
        author: Document author stored in core properties (optional).
        page_size: Page size — 'A4' or 'Letter' (default: A4).

    Returns:
        JSON with written file path and file size.
    """
    try:
        written = _md_to_docx(
            content, path,
            title or None,
            author or None,
            page_size,
        )
        size = Path(written).stat().st_size
        return json.dumps({
            "written_to": written,
            "file_size_bytes": size,
            "page_size": page_size,
        }, indent=2)

    except Exception as e:
        import traceback
        return json.dumps({"error": f"{type(e).__name__}: {e}\n{traceback.format_exc()}"})


# ──────────────────────────────────────────────────────────────────────────────
# /help HTML endpoint
# ──────────────────────────────────────────────────────────────────────────────

def help_handler(request: Request) -> HTMLResponse:
    raw_tools = list(mcp._tool_manager._tools.values())
    rows = ""
    for t in raw_tools:
        schema = getattr(t, "parameters", None) or {}
        props = schema.get("properties", {})
        required = set(schema.get("required", []))
        params = " ".join(
            f'<code class="param{" req" if k in required else ""}">{k}</code>'
            for k in props
        )
        rows += (
            f"<tr>"
            f"<td><b>{t.name}</b></td>"
            f"<td>{t.description or '<em>—</em>'}</td>"
            f"<td>{params or '<em>none</em>'}</td>"
            f"</tr>\n"
        )

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Docs MCP — Tools</title>
  <style>
    *{{box-sizing:border-box;margin:0;padding:0}}
    body{{font-family:'Segoe UI',system-ui,sans-serif;background:#0d1117;color:#c9d1d9;padding:2rem}}
    h1{{color:#d2a8ff;margin-bottom:.25rem;font-size:1.6rem}}
    p.sub{{color:#8b949e;margin-bottom:1.5rem;font-size:.9rem}}
    table{{width:100%;border-collapse:collapse;font-size:.9rem}}
    th{{background:#161b22;color:#8b949e;text-transform:uppercase;font-size:.75rem;
        letter-spacing:.05em;padding:.6rem 1rem;text-align:left;border-bottom:1px solid #30363d}}
    td{{padding:.65rem 1rem;border-bottom:1px solid #21262d;vertical-align:top}}
    tr:hover td{{background:#161b22}}
    b{{color:#e6edf3}}
    code.param{{background:#21262d;padding:.15em .4em;border-radius:4px;font-size:.82em;
               color:#79c0ff;margin-right:.25em}}
    code.param.req{{color:#ffa657}}
    em{{color:#6e7681}}
    .badge{{display:inline-block;background:#6e40c922;color:#d2a8ff;border:1px solid #8957e5;
            border-radius:12px;padding:.1rem .6rem;font-size:.75rem;margin-left:.5rem}}
    .note{{background:#161b22;border:1px solid #30363d;border-radius:6px;padding:1rem;
           margin-bottom:1.5rem;font-size:.85rem;color:#8b949e;line-height:1.6}}
    .note strong{{color:#c9d1d9}}
    .note code{{background:#21262d;padding:.1em .35em;border-radius:3px;color:#79c0ff;font-size:.85em}}
  </style>
</head>
<body>
  <h1>Docs MCP <span class="badge">{len(raw_tools)} tools</span></h1>
  <p class="sub">PDF reading and DOCX read/write for the AI assistant.</p>
  <div class="note">
    <strong>Transport:</strong> SSE &nbsp;·&nbsp;
    <strong>SSE endpoint:</strong> <code>/sse</code> &nbsp;·&nbsp;
    <strong>Port:</strong> <code>{PORT}</code>
    <br><br>
    Parameters in <code style="color:#ffa657">orange</code> are <strong>required</strong>.
  </div>
  <table>
    <thead><tr><th>Tool</th><th>Description</th><th>Parameters</th></tr></thead>
    <tbody>{rows}</tbody>
  </table>
</body>
</html>"""
    return HTMLResponse(html)


# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    from starlette.applications import Starlette
    from starlette.routing import Mount

    sse_starlette = mcp.sse_app()

    combined = Starlette(routes=[
        Route("/help", endpoint=help_handler),
        Mount("/", app=sse_starlette),
    ])

    uvicorn.run(combined, host="0.0.0.0", port=PORT, log_level="info")
